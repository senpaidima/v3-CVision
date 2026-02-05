from __future__ import annotations

import io
from pathlib import Path

import docx
import pytest
from starlette.testclient import TestClient

from app.core.dependencies import get_current_user
from app.main import app
from app.models.auth import UserInfo
from app.services.document_extractor import (
    MAX_FILE_SIZE,
    SUPPORTED_CONTENT_TYPES,
    DocumentExtractionError,
    DocumentExtractor,
)

FIXTURES_DIR = Path(__file__).resolve().parent.parent / "fixtures"


@pytest.fixture
def extractor():
    return DocumentExtractor()


@pytest.fixture
def sample_pdf_bytes():
    return (FIXTURES_DIR / "sample-lastenheft.pdf").read_bytes()


def _make_docx_bytes(text: str) -> bytes:
    doc = docx.Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_from_pdf_returns_text(extractor, sample_pdf_bytes):
    result = extractor.extract_from_pdf(sample_pdf_bytes)
    assert isinstance(result, str)
    assert len(result) > 0


def test_extract_from_pdf_invalid_bytes_raises(extractor):
    with pytest.raises(DocumentExtractionError, match="Failed to extract text from PDF"):
        extractor.extract_from_pdf(b"not a pdf")


def test_extract_from_docx_returns_text(extractor):
    docx_bytes = _make_docx_bytes("Anforderung: Das System soll Lebensläufe verwalten.")
    result = extractor.extract_from_docx(docx_bytes)
    assert "Anforderung" in result
    assert "Lebensläufe" in result


def test_extract_from_docx_invalid_bytes_raises(extractor):
    with pytest.raises(DocumentExtractionError, match="Failed to extract text from DOCX"):
        extractor.extract_from_docx(b"not a docx")


def test_extract_from_docx_empty_paragraphs_skipped(extractor):
    doc = docx.Document()
    doc.add_paragraph("Line one")
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("Line two")
    buf = io.BytesIO()
    doc.save(buf)
    result = extractor.extract_from_docx(buf.getvalue())
    assert result == "Line one\nLine two"


def test_extract_from_text_strips_whitespace(extractor):
    assert extractor.extract_from_text("  hello world  ") == "hello world"


def test_extract_from_text_empty_string(extractor):
    assert extractor.extract_from_text("   ") == ""


def test_extract_dispatches_pdf(extractor, sample_pdf_bytes):
    result = extractor.extract(sample_pdf_bytes, "application/pdf")
    assert isinstance(result, str)
    assert len(result) > 0


def test_extract_dispatches_docx(extractor):
    docx_bytes = _make_docx_bytes("Test content for DOCX dispatch")
    result = extractor.extract(
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "Test content" in result


def test_extract_rejects_unsupported_content_type(extractor):
    with pytest.raises(DocumentExtractionError, match="Unsupported content type"):
        extractor.extract(b"data", "application/octet-stream")


def test_extract_rejects_file_too_large(extractor):
    oversized = b"x" * (MAX_FILE_SIZE + 1)
    with pytest.raises(DocumentExtractionError, match="File too large"):
        extractor.extract(oversized, "application/pdf")


def test_extract_rejects_empty_file(extractor):
    with pytest.raises(DocumentExtractionError, match="Empty file"):
        extractor.extract(b"", "application/pdf")


@pytest.fixture
def auth_client():
    mock_user = UserInfo(id="test-1", name="Test User", email="test@emposo.de", roles=["admin"])
    app.dependency_overrides[get_current_user] = lambda: mock_user
    with TestClient(app) as c:
        yield c
    app.dependency_overrides.clear()


def test_upload_endpoint_returns_401_without_auth(client):
    response = client.post("/api/v1/lastenheft/upload")
    assert response.status_code == 401


def test_upload_endpoint_pdf_returns_extracted_text(auth_client, sample_pdf_bytes):
    response = auth_client.post(
        "/api/v1/lastenheft/upload",
        files={"file": ("test.pdf", sample_pdf_bytes, "application/pdf")},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["format"] == "pdf"
    assert data["char_count"] > 0
    assert len(data["extracted_text"]) == data["char_count"]


def test_upload_endpoint_rejects_invalid_type(auth_client):
    response = auth_client.post(
        "/api/v1/lastenheft/upload",
        files={"file": ("virus.exe", b"MZ\x90\x00", "application/octet-stream")},
    )
    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]


def test_upload_endpoint_rejects_empty_file(auth_client):
    response = auth_client.post(
        "/api/v1/lastenheft/upload",
        files={"file": ("empty.pdf", b"", "application/pdf")},
    )
    assert response.status_code == 400
    assert "empty" in response.json()["detail"].lower()


def test_text_paste_endpoint_returns_extracted_text(auth_client):
    response = auth_client.post(
        "/api/v1/lastenheft/text",
        json={"text": "  Das System soll CVs verwalten.  "},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["format"] == "text"
    assert data["extracted_text"] == "Das System soll CVs verwalten."
    assert data["char_count"] == 30


def test_text_paste_endpoint_returns_401_without_auth(client):
    response = client.post("/api/v1/lastenheft/text", json={"text": "test"})
    assert response.status_code == 401
